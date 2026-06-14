from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from report_data import build_rows, star_text

BRAND = "0B7285"
LIGHT = "E9F5F7"


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _hyperlink(paragraph, url, text):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def build_docx(data):
    doc = Document()

    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    _shade(cell, BRAND)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFINITE — Managed by MEDELITE")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("FACILITY ASSESSMENT SNAPSHOT")
    trun.bold = True
    trun.font.size = Pt(18)

    state = doc.add_paragraph()
    state.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = state.add_run(data.get("state", ""))
    srun.bold = True
    srun.font.size = Pt(12)
    srun.font.color.rgb = RGBColor.from_string(BRAND)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for kind, label, value in build_rows(data):
        row = table.add_row()
        if kind == "section":
            merged = row.cells[0].merge(row.cells[1])
            _shade(merged, BRAND)
            _set(merged, label, bold=True, color="FFFFFF", size=11)
        else:
            _shade(row.cells[0], LIGHT)
            _set(row.cells[0], label, bold=True)
            _set(row.cells[1], star_text(value) if kind == "rating" else str(value))

    link = data.get("source_url", "")
    src = doc.add_paragraph()
    src.add_run("Source: ")
    _hyperlink(src, link, link)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
